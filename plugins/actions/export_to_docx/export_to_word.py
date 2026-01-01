"""
title: Export to Word
author: Fu-Jie
author_url: https://github.com/Fu-Jie
funding_url: https://github.com/Fu-Jie/awesome-openwebui
version: 0.2.0
icon_url: data:image/svg+xml;base64,PHN2ZwogIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyIKICB3aWR0aD0iMjQiCiAgaGVpZ2h0PSIyNCIKICB2aWV3Qm94PSIwIDAgMjQgMjQiCiAgZmlsbD0ibm9uZSIKICBzdHJva2U9ImN1cnJlbnRDb2xvciIKICBzdHJva2Utd2lkdGg9IjIiCiAgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIgogIHN0cm9rZS1saW5lam9pbj0icm91bmQiCj4KICA8cGF0aCBkPSJNNiAyMmEyIDIgMCAwIDEtMi0yVjRhMiAyIDAgMCAxIDItMmg4YTIuNCAyLjQgMCAwIDEgMS43MDQuNzA2bDMuNTg4IDMuNTg4QTIuNCAyLjQgMCAwIDEgMjAgOHYxMmEyIDIgMCAwIDEtMiAyeiIgLz4KICA8cGF0aCBkPSJNMTQgMnY1YTEgMSAwIDAgMCAxIDFoNSIgLz4KICA8cGF0aCBkPSJNMTAgOUg4IiAvPgogIDxwYXRoIGQ9Ik0xNiAxM0g4IiAvPgogIDxwYXRoIGQ9Ik0xNiAxN0g4IiAvPgo8L3N2Zz4K
requirements: python-docx, httpx, Pygments, latex2mathml, mathml2omml
description: Export current conversation from Markdown to Word (.docx) with Mermaid (Kroki PNG), LaTeX math, real hyperlinks, improved tables, syntax highlighting, and blockquote support.
"""

import re
import base64
import datetime
import io
import asyncio
import logging
import hashlib
import ipaddress
import struct
import zlib
from dataclasses import dataclass
from typing import Optional, Callable, Awaitable, Any, List, Tuple, Dict, Literal, cast
from urllib.parse import urlparse

import httpx
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from open_webui.models.chats import Chats
from open_webui.models.users import Users
from open_webui.utils.chat import generate_chat_completion
from pydantic import BaseModel, Field

# Pygments for syntax highlighting
try:
    from pygments import lex
    from pygments.lexers import get_lexer_by_name, TextLexer
    from pygments.token import Token

    PYGMENTS_AVAILABLE = True
except ImportError:
    PYGMENTS_AVAILABLE = False

try:
    from latex2mathml.converter import convert as latex_to_mathml
    import mathml2omml

    LATEX_MATH_AVAILABLE = True
except Exception:
    LATEX_MATH_AVAILABLE = False


logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)

MermaidMode = Literal["off", "kroki"]
MermaidRendererSecurityMode = Literal["permissive", "strict"]
MermaidKrokiImageFormat = Literal["png", "svg", "svg+png"]

_AUTO_URL_RE = re.compile(r"(?:https?://|www\.)[^\s<>()]+")

_TRANSPARENT_1PX_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVQImWNgYGBgAAAABQABDQottAAAAABJRU5ErkJggg=="
)

_ASVG_NS = "http://schemas.microsoft.com/office/drawing/2016/SVG/main"
nsmap.setdefault("asvg", _ASVG_NS)

_REASONING_DETAILS_RE = re.compile(
    r"<details\b[^>]*\btype\s*=\s*(?:\"reasoning\"|'reasoning'|reasoning)[^>]*>.*?</details\s*>",
    re.IGNORECASE | re.DOTALL,
)
_THINK_RE = re.compile(r"<think\b[^>]*>.*?</think\s*>", re.IGNORECASE | re.DOTALL)
_ANALYSIS_RE = re.compile(
    r"<analysis\b[^>]*>.*?</analysis\s*>", re.IGNORECASE | re.DOTALL
)


class _MermaidResponseTooLarge(Exception):
    pass


@dataclass(frozen=True)
class _MermaidFenceBlock:
    info_raw: str
    language: str
    attrs: List[str]
    source: str


@dataclass(frozen=True)
class _MermaidOutcome:
    kind: Literal["image", "code"]
    png_bytes: Optional[bytes] = None
    svg_bytes: Optional[bytes] = None
    error_classification: Optional[str] = None
    error_detail: Optional[str] = None


@dataclass(frozen=True)
class _CitationRef:
    idx: int
    anchor: str
    title: str
    url: Optional[str]
    source_id: str


class Action:
    class Valves(BaseModel):
        TITLE_SOURCE: str = Field(
            default="chat_title",
            description="Title Source: 'chat_title' (Chat Title), 'ai_generated' (AI Generated), 'markdown_title' (Markdown Title)",
        )

        MERMAID_MODE: MermaidMode = Field(
            default="kroki",
            description="Mermaid conversion mode: off | kroki",
        )
        MERMAID_KROKI_BASE_URL: str = Field(
            default="https://kroki.io",
            description="Kroki base URL. Example: https://kroki.io or http://kroki:8000",
        )
        MERMAID_KROKI_TIMEOUT_S: int = Field(
            default=10, description="Kroki timeout (seconds)"
        )
        MERMAID_KROKI_MAX_REQUEST_BYTES: int = Field(
            default=100_000, description="Max Mermaid bytes per Kroki request"
        )
        MERMAID_KROKI_MAX_RESPONSE_BYTES: int = Field(
            default=8_000_000, description="Max PNG bytes per Kroki response"
        )
        MERMAID_KROKI_MAX_CONCURRENCY: int = Field(
            default=2, description="Max concurrent Kroki requests per export"
        )
        MERMAID_KROKI_IMAGE_FORMAT: MermaidKrokiImageFormat = Field(
            default="svg+png",
            description="Mermaid image format: png | svg | svg+png",
        )
        MERMAID_KROKI_BACKGROUND: str = Field(
            default="#FFFFFF",
            description="Mermaid background color for Kroki renders (empty disables)",
        )

        MERMAID_KROKI_SECURITY_MODE: MermaidRendererSecurityMode = Field(
            default="permissive",
            description="Kroki URL security: permissive | strict",
        )
        MERMAID_KROKI_ALLOWED_SCHEMES: str = Field(
            default="https",
            description="(strict) Allowed URL schemes, comma-separated",
        )
        MERMAID_KROKI_HOST_ALLOWLIST: str = Field(
            default="",
            description="(strict) Allowed renderer hostnames, comma-separated (empty = allow all)",
        )
        MERMAID_KROKI_BLOCK_PRIVATE_IPS: bool = Field(
            default=True, description="(strict) Block private IP targets"
        )
        MERMAID_KROKI_BLOCK_LOOPBACK: bool = Field(
            default=True, description="(strict) Block loopback targets"
        )

        MERMAID_MAX_TOTAL_IMAGE_BYTES_PER_EXPORT: int = Field(
            default=15_000_000,
            description="Max total Mermaid PNG bytes embedded per export (<=0 disables cap)",
        )

        MERMAID_CAPTIONS_ENABLE: bool = Field(
            default=True,
            description="Add figure captions under Mermaid images/charts",
        )
        MERMAID_CAPTION_STYLE: str = Field(
            default="Caption",
            description="Paragraph style name for Mermaid captions (uses 'Caption' if available, otherwise creates a safe custom style)",
        )
        MERMAID_CAPTION_PREFIX: str = Field(
            default="Figure",
            description="Caption prefix label, e.g. 'Figure'",
        )

        MATH_ENABLE: bool = Field(
            default=True,
            description="Enable LaTeX math block conversion (\\[...\\] and $$...$$) into Word equations",
        )

    def __init__(self):
        self.valves = self.Valves()
        self._mermaid_figure_counter: int = 0
        self._caption_style_name: Optional[str] = None
        self._citation_anchor_by_index: Dict[int, str] = {}
        self._citation_refs: List[_CitationRef] = []
        self._bookmark_id_counter: int = 1

    async def _send_notification(self, emitter: Callable, type: str, content: str):
        await emitter(
            {"type": "notification", "data": {"type": type, "content": content}}
        )

    async def action(
        self,
        body: dict,
        __user__=None,
        __event_emitter__=None,
        __event_call__: Optional[Callable[[Any], Awaitable[None]]] = None,
        __metadata__: Optional[dict] = None,
        __request__: Optional[Any] = None,
    ):
        logger.info(f"action:{__name__}")

        # Parse user info
        user_language = "en-US"
        user_name = "User"
        user_id = "unknown_user"
        if isinstance(__user__, (list, tuple)):
            user_language = (
                __user__[0].get("language", "en-US") if __user__ else "en-US"
            )
            user_name = __user__[0].get("name", "User") if __user__[0] else "User"
            user_id = (
                __user__[0]["id"]
                if __user__ and "id" in __user__[0]
                else "unknown_user"
            )
        elif isinstance(__user__, dict):
            user_language = __user__.get("language", "en-US")
            user_name = __user__.get("name", "User")
            user_id = __user__.get("id", "unknown_user")

        if __event_emitter__:
            last_assistant_message = body["messages"][-1]

            await __event_emitter__(
                {
                    "type": "status",
                    "data": {
                        "description": "Converting to Word document...",
                        "done": False,
                    },
                }
            )

            try:
                message_content = last_assistant_message["content"]
                if isinstance(message_content, str):
                    message_content = self._strip_reasoning_blocks(message_content)

                if not message_content or not message_content.strip():
                    await self._send_notification(
                        __event_emitter__, "error", "No content found to export!"
                    )
                    return

                # Generate filename
                title = ""
                chat_id = self.extract_chat_id(body, __metadata__)

                # Fetch chat_title directly via chat_id as it's usually missing in body
                chat_title = ""
                if chat_id:
                    chat_title = await self.fetch_chat_title(chat_id, user_id)

                if (
                    self.valves.TITLE_SOURCE.strip() == "chat_title"
                    or not self.valves.TITLE_SOURCE.strip()
                ):
                    title = chat_title
                elif self.valves.TITLE_SOURCE.strip() == "markdown_title":
                    title = self.extract_title(message_content)
                elif self.valves.TITLE_SOURCE.strip() == "ai_generated":
                    title = await self.generate_title_using_ai(
                        body, message_content, user_id, __request__
                    )

                # Fallback logic
                if not title:
                    if self.valves.TITLE_SOURCE.strip() != "chat_title" and chat_title:
                        title = chat_title
                    elif self.valves.TITLE_SOURCE.strip() != "markdown_title":
                        extracted = self.extract_title(message_content)
                        if extracted:
                            title = extracted

                current_datetime = datetime.datetime.now()
                formatted_date = current_datetime.strftime("%Y%m%d")

                if title:
                    filename = f"{self.clean_filename(title)}.docx"
                else:
                    filename = f"{user_name}_{formatted_date}.docx"

                top_heading = ""
                if chat_title:
                    top_heading = chat_title
                elif title:
                    top_heading = title

                # Create Word document; if no h1 exists, inject chat title as h1
                has_h1 = bool(re.search(r"^#\s+.+$", message_content, re.MULTILINE))
                sources = (
                    last_assistant_message.get("sources")
                    or body.get("sources")
                    or []
                )
                doc = await self.markdown_to_docx(
                    message_content,
                    top_heading=top_heading,
                    has_h1=has_h1,
                    sources=sources,
                )

                # Save to memory
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                file_content = doc_buffer.read()
                base64_blob = base64.b64encode(file_content).decode("utf-8")

                # Trigger file download
                if __event_call__:
                    await __event_call__(
                        {
                            "type": "execute",
                            "data": {
                                "code": f"""
                                try {{
                                    const base64Data = "{base64_blob}";
                                    const binaryData = atob(base64Data);
                                    const arrayBuffer = new Uint8Array(binaryData.length);
                                    for (let i = 0; i < binaryData.length; i++) {{
                                        arrayBuffer[i] = binaryData.charCodeAt(i);
                                    }}
                                    const blob = new Blob([arrayBuffer], {{ type: "application/vnd.openxmlformats-officedocument.wordprocessingml.document" }});
                                    const filename = "{filename}";

                                    const url = URL.createObjectURL(blob);
                                    const a = document.createElement("a");
                                    a.style.display = "none";
                                    a.href = url;
                                    a.download = filename;
                                    document.body.appendChild(a);
                                    a.click();
                                    URL.revokeObjectURL(url);
                                    document.body.removeChild(a);
                                }} catch (error) {{
                                    console.error('Error triggering download:', error);
                                }}
                                """
                            },
                        }
                    )

                await __event_emitter__(
                    {
                        "type": "status",
                        "data": {"description": "Word document exported", "done": True},
                    }
                )

                await self._send_notification(
                    __event_emitter__, "success", f"Successfully exported to {filename}"
                )

                return {"message": "Download triggered"}

            except Exception as e:
                logger.exception(f"Error exporting to Word: {str(e)}")
                await __event_emitter__(
                    {
                        "type": "status",
                        "data": {
                            "description": f"Export failed: {str(e)}",
                            "done": True,
                        },
                    }
                )
                await self._send_notification(
                    __event_emitter__,
                    "error",
                    f"Error exporting Word document: {str(e)}",
                )

    async def generate_title_using_ai(
        self, body: dict, content: str, user_id: str, request: Any
    ) -> str:
        if not request:
            return ""

        try:
            user_obj = Users.get_user_by_id(user_id)
            model = body.get("model")

            payload = {
                "model": model,
                "messages": [
                    {
                        "role": "system",
                        "content": "You are a helpful assistant. Generate a short, concise title (max 10 words) for the following text. Do not use quotes. Only output the title.",
                    },
                    {"role": "user", "content": content[:2000]},  # Limit content length
                ],
                "stream": False,
            }

            response = await generate_chat_completion(request, payload, user_obj)
            if response and "choices" in response:
                return response["choices"][0]["message"]["content"].strip()
        except Exception as e:
            logger.error(f"Error generating title: {e}")

        return ""

    def extract_title(self, content: str) -> str:
        """Extract title from Markdown h1/h2 only"""
        lines = content.split("\n")
        for line in lines:
            # Match h1-h2 headings only
            match = re.match(r"^#{1,2}\s+(.+)$", line.strip())
            if match:
                return match.group(1).strip()
        return ""

    def extract_chat_title(self, body: dict) -> str:
        """Extract chat title from common payload fields."""
        if not isinstance(body, dict):
            return ""

        candidates = []

        for key in ("chat", "conversation"):
            if isinstance(body.get(key), dict):
                candidates.append(body.get(key, {}).get("title", ""))

        for key in ("title", "chat_title"):
            value = body.get(key)
            if isinstance(value, str):
                candidates.append(value)

        for candidate in candidates:
            if candidate and isinstance(candidate, str):
                return candidate.strip()
        return ""

    def extract_chat_id(self, body: dict, metadata: Optional[dict]) -> str:
        """Extract chat_id from body or metadata"""
        if isinstance(body, dict):
            chat_id = body.get("chat_id") or body.get("id")
            if isinstance(chat_id, str) and chat_id.strip():
                return chat_id.strip()

            for key in ("chat", "conversation"):
                nested = body.get(key)
                if isinstance(nested, dict):
                    nested_id = nested.get("id") or nested.get("chat_id")
                    if isinstance(nested_id, str) and nested_id.strip():
                        return nested_id.strip()
        if isinstance(metadata, dict):
            chat_id = metadata.get("chat_id")
            if isinstance(chat_id, str) and chat_id.strip():
                return chat_id.strip()
        return ""

    async def fetch_chat_title(self, chat_id: str, user_id: str = "") -> str:
        """Fetch chat title from database by chat_id"""
        if not chat_id:
            return ""

        def _load_chat():
            if user_id:
                return Chats.get_chat_by_id_and_user_id(id=chat_id, user_id=user_id)
            return Chats.get_chat_by_id(chat_id)

        try:
            chat = await asyncio.to_thread(_load_chat)
        except Exception as exc:
            logger.warning(f"Failed to load chat {chat_id}: {exc}")
            return ""

        if not chat:
            return ""

        data = getattr(chat, "chat", {}) or {}
        title = data.get("title") or getattr(chat, "title", "")
        return title.strip() if isinstance(title, str) else ""

    def clean_filename(self, name: str) -> str:
        """Clean illegal characters from filename"""
        return re.sub(r'[\\/*?:"<>|]', "", name).strip()[:50]

    async def markdown_to_docx(
        self,
        markdown_text: str,
        top_heading: str = "",
        has_h1: bool = False,
        sources: Optional[List[dict]] = None,
    ) -> Document:
        """
        Convert Markdown text to Word document
        Supports: headings, paragraphs, bold, italic, code blocks, lists, tables, links
        Additionally: Mermaid fenced blocks (```mermaid) rendered via Kroki (PNG/SVG),
        LaTeX math to Word equations, and OpenWebUI citations to References.
        """
        doc = Document()
        self._mermaid_figure_counter = 0
        self._caption_style_name = None
        self._citation_anchor_by_index = {}
        self._citation_refs = self._build_citation_refs(sources or [])
        self._bookmark_id_counter = 1
        for ref in self._citation_refs:
            self._citation_anchor_by_index[ref.idx] = ref.anchor

        # Set default fonts
        self.set_document_default_font(doc)

        # If there is no h1 in content, prepend chat title as h1 when provided
        if top_heading and not has_h1:
            self.add_heading(doc, top_heading, 1)

        mermaid_outcomes = await self._precompute_mermaid_outcomes(markdown_text)
        mermaid_outcome_cursor = 0

        lines = markdown_text.split("\n")
        i = 0
        in_code_block = False
        code_block_content = []
        code_block_info_raw = ""
        code_block_lang = ""
        code_block_attrs: List[str] = []
        in_math_block = False
        math_block_delim = ""
        math_block_lines: List[str] = []
        in_list = False
        list_items = []
        list_type = None  # 'ordered' or 'unordered'

        while i < len(lines):
            line = lines[i]

            # Handle display math blocks (\[...\] or $$...$$)
            if not in_code_block and self.valves.MATH_ENABLE:
                single_line = self._extract_single_line_math(line)
                if single_line is not None:
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                        in_list = False
                    self._add_display_equation(doc, single_line)
                    i += 1
                    continue

                if not in_math_block:
                    stripped = line.strip()
                    if stripped in (r"\[", "$$"):
                        if in_list and list_items:
                            self.add_list_to_doc(doc, list_items, list_type)
                            list_items = []
                            in_list = False
                        in_math_block = True
                        math_block_delim = stripped
                        math_block_lines = []
                        i += 1
                        continue
                else:
                    stripped = line.strip()
                    close = r"\]" if math_block_delim == r"\[" else "$$"
                    if stripped == close:
                        in_math_block = False
                        latex = "\n".join(math_block_lines).strip()
                        self._add_display_equation(doc, latex)
                        math_block_delim = ""
                        math_block_lines = []
                        i += 1
                        continue
                    math_block_lines.append(line)
                    i += 1
                    continue

            # Handle code blocks
            if line.strip().startswith("```"):
                if not in_code_block:
                    # Process pending list first
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                        in_list = False

                    in_code_block = True
                    code_block_info_raw = line.strip()[3:].strip()
                    code_block_lang, code_block_attrs = self._parse_fence_info(
                        code_block_info_raw
                    )
                    code_block_content = []
                else:
                    # End code block
                    in_code_block = False
                    code_text = "\n".join(code_block_content)
                    if code_block_lang.lower() == "mermaid":
                        outcome = (
                            mermaid_outcomes[mermaid_outcome_cursor]
                            if mermaid_outcome_cursor < len(mermaid_outcomes)
                            else _MermaidOutcome(
                                kind="code",
                                error_classification="mermaid_preprocessing_mismatch",
                                error_detail="Mermaid preprocessing mismatch",
                            )
                        )
                        mermaid_outcome_cursor += 1
                        self._insert_mermaid_outcome(
                            doc,
                            code_text,
                            outcome,
                            info_raw=code_block_info_raw,
                            attrs=code_block_attrs,
                        )
                    else:
                        self.add_code_block(doc, code_text, code_block_lang)
                    code_block_content = []
                    code_block_info_raw = ""
                    code_block_lang = ""
                    code_block_attrs = []
                i += 1
                continue

            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue

            # Handle tables
            if line.strip().startswith("|") and line.strip().endswith("|"):
                # Process pending list first
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False

                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                self.add_table(doc, table_lines)
                continue

            # Handle headings
            header_match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
            if header_match:
                # Process pending list first
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False

                level = len(header_match.group(1))
                text = header_match.group(2)
                self.add_heading(doc, text, level)
                i += 1
                continue

            # Handle unordered lists
            unordered_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
            if unordered_match:
                if not in_list or list_type != "unordered":
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                    in_list = True
                    list_type = "unordered"
                indent = len(unordered_match.group(1)) // 2
                list_items.append((indent, unordered_match.group(2)))
                i += 1
                continue

            # Handle ordered lists
            ordered_match = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
            if ordered_match:
                if not in_list or list_type != "ordered":
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                    in_list = True
                    list_type = "ordered"
                indent = len(ordered_match.group(1)) // 2
                list_items.append((indent, ordered_match.group(2)))
                i += 1
                continue

            # Handle blockquotes
            if line.strip().startswith(">"):
                # Process pending list first
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False

                # Collect consecutive quote lines
                blockquote_lines = []
                while i < len(lines) and lines[i].strip().startswith(">"):
                    # Remove leading > and optional space
                    quote_line = re.sub(r"^>\s?", "", lines[i])
                    blockquote_lines.append(quote_line)
                    i += 1
                self.add_blockquote(doc, "\n".join(blockquote_lines))
                continue

            # Handle horizontal rules
            if re.match(r"^[-*_]{3,}$", line.strip()):
                # Process pending list first
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False

                self.add_horizontal_rule(doc)
                i += 1
                continue

            # Handle empty lines
            if not line.strip():
                # End list
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False
                i += 1
                continue

            # Handle normal paragraphs
            if in_list and list_items:
                self.add_list_to_doc(doc, list_items, list_type)
                list_items = []
                in_list = False

            self.add_paragraph(doc, line)
            i += 1

        # Process remaining list
        if in_list and list_items:
            self.add_list_to_doc(doc, list_items, list_type)

        # If math block wasn't closed, render it as plain text for robustness.
        if in_math_block and math_block_lines:
            self.add_paragraph(doc, r"\[")
            for l in math_block_lines:
                self.add_paragraph(doc, l)
            self.add_paragraph(doc, r"\]")

        if self._citation_refs:
            self._add_references_section(doc)

        return doc

    def _extract_single_line_math(self, line: str) -> Optional[str]:
        s = line.strip()
        # \[ ... \]
        m = re.match(r"^\\\[(.*)\\\]$", s)
        if m:
            return m.group(1).strip()
        # $$ ... $$
        m = re.match(r"^\$\$(.*)\$\$$", s)
        if m:
            return m.group(1).strip()
        return None

    def _strip_reasoning_blocks(self, text: str) -> str:
        """
        Strip model reasoning blocks from assistant Markdown before export.

        OpenWebUI can include reasoning as interleaved <details type=\"reasoning\">...</details>
        (and sometimes <think>/<analysis> blocks). These should never be exported into DOCX.
        """
        if not text:
            return text

        cur = text
        for _ in range(10):
            prev = cur
            cur = _REASONING_DETAILS_RE.sub("", cur)
            cur = _THINK_RE.sub("", cur)
            cur = _ANALYSIS_RE.sub("", cur)
            if cur == prev:
                break

        # Clean up excessive blank lines left by removals.
        cur = re.sub(r"\n{4,}", "\n\n\n", cur)
        return cur

    def _add_display_equation(self, doc: Document, latex: str):
        latex = (latex or "").strip()
        if not latex:
            return

        if not LATEX_MATH_AVAILABLE:
            self.add_code_block(doc, latex, "latex")
            return

        try:
            mathml = latex_to_mathml(latex)
            omml = mathml2omml.convert(mathml)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cast(Any, para)._p.append(self._wrap_omml_for_word(omml))
        except Exception as exc:
            logger.warning(f"Math conversion failed; falling back to text: {exc}")
            self.add_code_block(doc, latex, "latex")

    def _wrap_omml_for_word(self, omml: str):
        m_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        # Keep the OMML payload as-is, but ensure it has the math namespace declared.
        xml = f'<m:oMathPara xmlns:m="{m_ns}" xmlns:w="{w_ns}">{omml}</m:oMathPara>'
        return parse_xml(xml)

    # (Math warning paragraphs removed)

    def _build_citation_refs(self, sources: List[dict]) -> List[_CitationRef]:
        citation_idx_map: Dict[str, int] = {}
        refs_by_idx: Dict[int, _CitationRef] = {}

        for source in sources or []:
            if not isinstance(source, dict):
                continue

            documents = source.get("document") or []
            metadatas = source.get("metadata") or []
            src_info = source.get("source") or {}

            src_name = src_info.get("name") if isinstance(src_info, dict) else None
            src_id_default = (
                src_info.get("id") if isinstance(src_info, dict) else None
            )
            src_urls = (
                src_info.get("urls") if isinstance(src_info, dict) else None
            )

            if not isinstance(documents, list):
                documents = []
            if not isinstance(metadatas, list):
                metadatas = []

            for idx_doc, _doc_text in enumerate(documents):
                meta = metadatas[idx_doc] if idx_doc < len(metadatas) else {}
                if not isinstance(meta, dict):
                    meta = {}

                source_id = (
                    meta.get("source")
                    or src_id_default
                    or "N/A"
                )
                source_id_str = str(source_id)

                if source_id_str not in citation_idx_map:
                    citation_idx_map[source_id_str] = len(citation_idx_map) + 1
                idx = citation_idx_map[source_id_str]

                if idx in refs_by_idx:
                    continue

                url: Optional[str] = None
                if isinstance(source_id, str) and re.match(r"^https?://", source_id):
                    url = source_id
                elif isinstance(meta.get("url"), str) and re.match(r"^https?://", meta["url"]):
                    url = meta["url"]
                elif isinstance(src_urls, list) and src_urls:
                    if isinstance(src_urls[0], str) and re.match(r"^https?://", src_urls[0]):
                        url = src_urls[0]

                title = (
                    (meta.get("title") if isinstance(meta.get("title"), str) else None)
                    or (meta.get("name") if isinstance(meta.get("name"), str) else None)
                    or (src_name if isinstance(src_name, str) and src_name.strip() else None)
                    or (url if url else None)
                    or source_id_str
                )

                anchor = f"OWUIRef{idx}"
                refs_by_idx[idx] = _CitationRef(
                    idx=idx,
                    anchor=anchor,
                    title=title,
                    url=url,
                    source_id=source_id_str,
                )

        return [refs_by_idx[i] for i in sorted(refs_by_idx.keys())]

    def _add_bookmark(self, paragraph, name: str):
        bookmark_id = self._bookmark_id_counter
        self._bookmark_id_counter += 1

        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), name)

        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))

        p = cast(Any, paragraph)._p
        p.insert(0, start)
        p.append(end)

    def _add_internal_hyperlink(self, paragraph, display_text: str, anchor: str):
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), anchor)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)

        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = display_text
        new_run.append(t)

        hyperlink.append(new_run)
        cast(Any, paragraph)._p.append(hyperlink)

    def _add_references_section(self, doc: Document):
        self.add_heading(doc, "References", 2)

        for ref in self._citation_refs:
            para = doc.add_paragraph(style="List Number")
            self._add_bookmark(para, ref.anchor)
            # Include URL as an external link when available.
            if ref.url:
                self._add_hyperlink(para, ref.title, ref.url, display_text=ref.title)
            else:
                self._add_text_run(para, ref.title, bold=False, italic=False, strike=False)

    def _parse_fence_info(self, info_raw: str) -> Tuple[str, List[str]]:
        parts = [p for p in (info_raw or "").split() if p.strip()]
        if not parts:
            return "", []
        return parts[0], parts[1:]

    def _extract_mermaid_fences(self, markdown_text: str) -> List[_MermaidFenceBlock]:
        lines = markdown_text.split("\n")
        in_code_block = False
        info_raw = ""
        lang = ""
        attrs: List[str] = []
        content: List[str] = []
        blocks: List[_MermaidFenceBlock] = []

        for line in lines:
            if line.strip().startswith("```"):
                if not in_code_block:
                    in_code_block = True
                    info_raw = line.strip()[3:].strip()
                    lang, attrs = self._parse_fence_info(info_raw)
                    content = []
                else:
                    in_code_block = False
                    if lang.lower() == "mermaid":
                        blocks.append(
                            _MermaidFenceBlock(
                                info_raw=info_raw,
                                language=lang,
                                attrs=list(attrs),
                                source="\n".join(content),
                            )
                        )
                    info_raw = ""
                    lang = ""
                    attrs = []
                    content = []
                continue

            if in_code_block:
                content.append(line)

        return blocks

    def _normalize_mermaid_text(self, source: str) -> str:
        text = (source or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        return text + "\n"

    def _prepare_mermaid_for_kroki(self, source: str) -> str:
        text = self._strip_mermaid_title_for_render(source)
        bg = (self.valves.MERMAID_KROKI_BACKGROUND or "").strip()
        if bg == "":
            return text
        escaped = bg.replace("\\", "\\\\").replace('"', '\\"')
        init_line = f'%%{{init: {{"themeVariables": {{"background": "{escaped}"}}}}}}%%'
        # Avoid duplicating if already present (idempotent).
        if init_line in text:
            return text
        return init_line + "\n" + text

    def _mermaid_cache_key(self, source: str) -> str:
        normalized = self._normalize_mermaid_text(source)
        kroki_base = (self.valves.MERMAID_KROKI_BASE_URL or "").strip()
        key_material = "\n".join(
            [
                "v1",
                f"mode={self.valves.MERMAID_MODE}",
                f"format={self.valves.MERMAID_KROKI_IMAGE_FORMAT}",
                f"kroki={kroki_base}",
                f"timeout={self.valves.MERMAID_KROKI_TIMEOUT_S}",
                f"max_req={self.valves.MERMAID_KROKI_MAX_REQUEST_BYTES}",
                f"max_resp={self.valves.MERMAID_KROKI_MAX_RESPONSE_BYTES}",
                normalized,
            ]
        ).encode("utf-8", errors="replace")
        return hashlib.sha256(key_material).hexdigest()

    async def _precompute_mermaid_outcomes(
        self, markdown_text: str
    ) -> List[_MermaidOutcome]:
        blocks = self._extract_mermaid_fences(markdown_text)
        if not blocks:
            return []

        mode = self.valves.MERMAID_MODE

        total_budget_cap = self.valves.MERMAID_MAX_TOTAL_IMAGE_BYTES_PER_EXPORT
        budget_enabled = total_budget_cap > 0
        budget_used = 0
        budget_exceeded = False

        cache: Dict[str, _MermaidOutcome] = {}

        renderer_ok = await self._validate_renderer_base_url()
        can_render = mode == "kroki" and renderer_ok

        concurrency = self.valves.MERMAID_KROKI_MAX_CONCURRENCY
        semaphore = asyncio.Semaphore(concurrency)

        outcomes: List[_MermaidOutcome] = []

        async with httpx.AsyncClient(follow_redirects=False) as client:
            idx = 0
            while idx < len(blocks):
                block = blocks[idx]
                cache_key = self._mermaid_cache_key(block.source)

                cached = cache.get(cache_key)
                if cached is not None:
                    adjusted, budget_used, budget_exceeded = self._apply_image_budget(
                        cached, budget_enabled, total_budget_cap, budget_used, budget_exceeded
                    )
                    outcomes.append(adjusted)
                    idx += 1
                    continue

                if mode == "off":
                    outcome = _MermaidOutcome(
                        kind="code",
                        error_classification="renderer_disabled",
                        error_detail="MERMAID_MODE=off",
                    )
                    cache[cache_key] = outcome
                    outcomes.append(outcome)
                    idx += 1
                    continue

                if not can_render:
                    outcome = _MermaidOutcome(
                        kind="code",
                        error_classification="renderer_disabled",
                        error_detail="Renderer not enabled/available",
                    )
                    cache[cache_key] = outcome
                    outcomes.append(outcome)
                    idx += 1
                    continue

                if budget_exceeded and budget_enabled:
                    outcome = _MermaidOutcome(
                        kind="code",
                        error_classification="image_budget_exceeded",
                        error_detail="Max total Mermaid image bytes reached",
                    )
                    cache[cache_key] = outcome
                    outcomes.append(outcome)
                    idx += 1
                    continue

                # Batch up to `concurrency` renderer calls for better throughput.
                batch: List[Tuple[int, str, _MermaidFenceBlock]] = []
                while (
                    idx < len(blocks)
                    and len(batch) < concurrency
                    and not (budget_exceeded and budget_enabled)
                ):
                    blk = blocks[idx]
                    key = self._mermaid_cache_key(blk.source)
                    if key in cache:
                        adjusted, budget_used, budget_exceeded = self._apply_image_budget(
                            cache[key],
                            budget_enabled,
                            total_budget_cap,
                            budget_used,
                            budget_exceeded,
                        )
                        outcomes.append(adjusted)
                        idx += 1
                        continue

                    if budget_exceeded and budget_enabled:
                        outcome = _MermaidOutcome(
                            kind="code",
                            error_classification="image_budget_exceeded",
                            error_detail="Max total Mermaid image bytes reached",
                        )
                        cache[key] = outcome
                        outcomes.append(outcome)
                        idx += 1
                        continue

                    batch.append((idx, key, blk))
                    idx += 1

                if not batch:
                    continue

                # Deduplicate identical blocks within the batch.
                unique: Dict[str, _MermaidFenceBlock] = {}
                for _, key, blk in batch:
                    if key not in unique:
                        unique[key] = blk

                async def _render_one(key: str, blk: _MermaidFenceBlock):
                    async with semaphore:
                        outcome = await self._render_mermaid_for_export(
                            client, blk.source
                        )
                        cache[key] = outcome

                await asyncio.gather(*[_render_one(k, b) for k, b in unique.items()])

                # Apply outcomes in order, enforcing the image budget.
                for _, key, _ in batch:
                    out = cache.get(key) or _MermaidOutcome(
                        kind="code",
                        error_classification="renderer_network_error",
                        error_detail="Renderer task missing result",
                    )
                    adjusted, budget_used, budget_exceeded = self._apply_image_budget(
                        out, budget_enabled, total_budget_cap, budget_used, budget_exceeded
                    )
                    outcomes.append(adjusted)

        return outcomes

    def _apply_image_budget(
        self,
        outcome: _MermaidOutcome,
        budget_enabled: bool,
        budget_cap: int,
        budget_used: int,
        budget_exceeded: bool,
    ) -> Tuple[_MermaidOutcome, int, bool]:
        if outcome.kind != "image":
            return outcome, budget_used, budget_exceeded

        if not budget_enabled:
            return outcome, budget_used, budget_exceeded

        if budget_exceeded:
            return (
                _MermaidOutcome(
                    kind="code",
                    error_classification="image_budget_exceeded",
                    error_detail="Max total Mermaid image bytes reached",
                ),
                budget_used,
                True,
            )

        image_len = 0
        if outcome.png_bytes:
            image_len += len(outcome.png_bytes)
        if outcome.svg_bytes:
            image_len += len(outcome.svg_bytes)

        if image_len == 0:
            return outcome, budget_used, budget_exceeded

        if budget_used + image_len > budget_cap:
            return (
                _MermaidOutcome(
                    kind="code",
                    error_classification="image_budget_exceeded",
                    error_detail=f"Budget cap {budget_cap} bytes exceeded",
                ),
                budget_used,
                True,
            )

        return outcome, budget_used + image_len, False

    async def _render_mermaid_for_export(
        self, client: httpx.AsyncClient, source: str
    ) -> _MermaidOutcome:
        fmt = self.valves.MERMAID_KROKI_IMAGE_FORMAT
        if fmt == "png":
            return await self._render_mermaid_to_png(client, source)
        if fmt == "svg":
            svg_out = await self._render_mermaid_to_svg(client, source)
            if svg_out.kind != "image" or not svg_out.svg_bytes:
                return svg_out
            # In svg-only mode, embed a transparent PNG fallback with the same aspect ratio.
            # Word uses the fallback PNG's intrinsic ratio to size the drawing; a 1x1 fallback
            # makes wide charts look squashed.
            fallback_png = self._transparent_png_for_svg(svg_out.svg_bytes)
            return _MermaidOutcome(
                kind="image",
                png_bytes=fallback_png,
                svg_bytes=svg_out.svg_bytes,
            )

        # svg+png
        png_out = await self._render_mermaid_to_png(client, source)
        if png_out.kind != "image" or not png_out.png_bytes:
            return png_out
        svg_out = await self._render_mermaid_to_svg(client, source)
        if svg_out.kind == "image" and svg_out.svg_bytes:
            return _MermaidOutcome(
                kind="image",
                png_bytes=png_out.png_bytes,
                svg_bytes=svg_out.svg_bytes,
            )
        return png_out

    async def _validate_renderer_base_url(self) -> bool:
        base = (self.valves.MERMAID_KROKI_BASE_URL or "").strip()
        if not base:
            logger.warning("MERMAID_KROKI_BASE_URL is empty; Kroki disabled")
            return False

        parsed = urlparse(base)
        if not parsed.scheme or not parsed.netloc:
            logger.warning(f"Invalid MERMAID_KROKI_BASE_URL={base!r}; Kroki disabled")
            return False

        if self.valves.MERMAID_KROKI_SECURITY_MODE == "permissive":
            return True

        allowed_schemes_raw = self.valves.MERMAID_KROKI_ALLOWED_SCHEMES or ""
        allowed_schemes = {
            s.strip().lower() for s in allowed_schemes_raw.split(",") if s.strip()
        }
        if allowed_schemes and parsed.scheme.lower() not in allowed_schemes:
            logger.warning(f"Renderer scheme not allowed: {parsed.scheme!r}")
            return False

        host = parsed.hostname or ""
        allowlist_raw = self.valves.MERMAID_KROKI_HOST_ALLOWLIST or ""
        allowlist = {h.strip().lower() for h in allowlist_raw.split(",") if h.strip()}
        if allowlist and host.lower() not in allowlist:
            logger.warning(f"Renderer host not in allowlist: {host!r}")
            return False

        try:
            addrs = await asyncio.wait_for(
                asyncio.get_running_loop().getaddrinfo(host, parsed.port or 443),
                timeout=2.0,
            )
        except Exception as exc:
            logger.warning(f"Failed to resolve renderer host {host!r}: {exc}")
            return False

        block_private = self.valves.MERMAID_KROKI_BLOCK_PRIVATE_IPS
        block_loopback = self.valves.MERMAID_KROKI_BLOCK_LOOPBACK
        for family, _, _, _, sockaddr in addrs:
            ip_str = sockaddr[0]
            try:
                ip = ipaddress.ip_address(ip_str)
            except Exception:
                continue
            if block_loopback and ip.is_loopback:
                logger.warning(f"Renderer target resolves to loopback: {ip_str}")
                return False
            if block_private and (ip.is_private or ip.is_link_local or ip.is_reserved):
                logger.warning(f"Renderer target resolves to private/reserved IP: {ip_str}")
                return False

        return True

    async def _render_mermaid_to_png(
        self, client: httpx.AsyncClient, source: str
    ) -> _MermaidOutcome:
        mermaid_text = self._prepare_mermaid_for_kroki(source)
        data = mermaid_text.encode("utf-8", errors="replace")
        if len(data) > self.valves.MERMAID_KROKI_MAX_REQUEST_BYTES:
            return _MermaidOutcome(
                kind="code",
                error_classification="renderer_request_too_large",
                error_detail=f"{len(data)} bytes > {self.valves.MERMAID_KROKI_MAX_REQUEST_BYTES}",
            )

        base = (self.valves.MERMAID_KROKI_BASE_URL or "").strip().rstrip("/")
        url = f"{base}/mermaid/png"
        timeout = httpx.Timeout(self.valves.MERMAID_KROKI_TIMEOUT_S)

        try:
            async with client.stream(
                "POST",
                url,
                content=data,
                headers={"Content-Type": "text/plain", "Accept": "image/png"},
                timeout=timeout,
            ) as resp:
                status = resp.status_code
                content_type = (resp.headers.get("content-type") or "").lower()

                if status != 200:
                    snippet = await self._read_text_snippet(resp, limit=300)
                    classification = "renderer_http_error"
                    if status in (400, 422):
                        classification = "renderer_http_syntax_error"
                    elif status == 429:
                        classification = "renderer_http_rate_limited"
                    elif 500 <= status <= 599:
                        classification = "renderer_http_server_error"

                    logger.warning(
                        f"Mermaid renderer non-200: status={status} content-type={content_type!r} body={snippet!r}"
                    )
                    return _MermaidOutcome(
                        kind="code",
                        error_classification=classification,
                        error_detail=str(status),
                    )

                png = await self._read_limited_bytes(
                    resp, limit=self.valves.MERMAID_KROKI_MAX_RESPONSE_BYTES
                )
                if not self._looks_like_png(png, content_type):
                    snippet = png[:300].decode("utf-8", errors="replace")
                    logger.warning(
                        f"Mermaid renderer returned unexpected payload: content-type={content_type!r} snippet={snippet!r}"
                    )
                    return _MermaidOutcome(
                        kind="code",
                        error_classification="renderer_unexpected_payload",
                        error_detail=content_type or "unknown",
                    )

                return _MermaidOutcome(kind="image", png_bytes=png)

        except httpx.TimeoutException:
            return _MermaidOutcome(kind="code", error_classification="renderer_timeout")
        except httpx.RequestError as exc:
            logger.warning(f"Mermaid renderer network error: {exc}")
            return _MermaidOutcome(kind="code", error_classification="renderer_network_error")
        except _MermaidResponseTooLarge:
            return _MermaidOutcome(
                kind="code", error_classification="renderer_response_too_large"
            )
        except Exception as exc:
            logger.exception(f"Mermaid renderer unexpected error: {exc}")
            return _MermaidOutcome(kind="code", error_classification="renderer_network_error")

    def _looks_like_svg(self, data: bytes, content_type: str) -> bool:
        if content_type and "svg" in content_type:
            return True
        head = data.lstrip()[:1024].lower()
        return head.startswith(b"<svg") or b"<svg" in head

    def _pad_svg_viewbox(self, svg_bytes: bytes) -> bytes:
        """
        Kroki/Mermaid SVGs can have tight viewBoxes that clip strokes/text near edges
        in Word's SVG renderer. Expand the viewBox slightly to add padding.
        """
        try:
            s = svg_bytes.decode("utf-8", errors="strict")
        except Exception:
            return svg_bytes

        # First viewBox only (on the root <svg>).
        m = re.search(
            r'viewBox="([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)\s+'
            r'([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)\s+'
            r'([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)\s+'
            r'([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)"',
            s,
        )
        if not m:
            return svg_bytes

        try:
            min_x = float(m.group(1))
            min_y = float(m.group(2))
            width = float(m.group(3))
            height = float(m.group(4))
        except Exception:
            return svg_bytes

        if width <= 0 or height <= 0:
            return svg_bytes

        # Keep padding conservative; Word dark-mode “missing text” is often contrast, not clip.
        min_dim = min(width, height)
        pad = max(8.0, min_dim * 0.02)
        pad = min(pad, 24.0)

        new_vb = (
            f'viewBox="{min_x - pad:g} {min_y - pad:g} '
            f"{width + (2 * pad):g} {height + (2 * pad):g}\""
        )
        s2 = s[: m.start()] + new_vb + s[m.end() :]

        return s2.encode("utf-8", errors="replace")

    def _normalize_svg_for_word(self, svg_bytes: bytes) -> bytes:
        """
        Word's SVG renderer is sensitive to percentage sizing and certain CSS.
        Normalize the root <svg> to an explicit width/height (from viewBox) and
        remove max-width styling that can contribute to clipping/squashing.
        """
        try:
            s = svg_bytes.decode("utf-8", errors="strict")
        except Exception:
            return svg_bytes

        msvg = re.search(r"<svg\b[^>]*>", s, re.IGNORECASE)
        if not msvg:
            return svg_bytes

        tag = msvg.group(0)

        mvb = re.search(
            r'viewBox="([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)\s+'
            r'([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)\s+'
            r'([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)\s+'
            r'([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)"',
            tag,
        )
        vbw = vbh = None
        if mvb:
            try:
                vbw = float(mvb.group(3))
                vbh = float(mvb.group(4))
            except Exception:
                vbw = vbh = None

        def _get_attr(attr: str) -> Optional[str]:
            mm = re.search(rf'\b{attr}="([^"]*)"', tag, re.IGNORECASE)
            return mm.group(1) if mm else None

        width_attr = _get_attr("width")
        height_attr = _get_attr("height")

        needs_explicit = (
            vbw
            and vbh
            and (
                not width_attr
                or not height_attr
                or width_attr.strip().endswith("%")
                or height_attr.strip().endswith("%")
            )
        )

        tag2 = tag
        if needs_explicit:
            tag2 = re.sub(r'\bwidth="[^"]*"', "", tag2, flags=re.IGNORECASE)
            tag2 = re.sub(r'\bheight="[^"]*"', "", tag2, flags=re.IGNORECASE)
            tag2 = tag2[:-1] + f' width="{vbw:g}" height="{vbh:g}">'

        # Remove inline style max-width which can confuse Word sizing.
        tag2 = re.sub(r'\s+style="[^"]*"', "", tag2, flags=re.IGNORECASE)

        if "preserveAspectRatio=" not in tag2:
            tag2 = tag2[:-1] + ' preserveAspectRatio="xMidYMid meet">'

        if "overflow=" not in tag2.lower():
            tag2 = tag2[:-1] + ' overflow="visible">'

        if tag2 == tag:
            s2 = s
        else:
            s2 = s[: msvg.start()] + tag2 + s[msvg.end() :]

        bg = (self.valves.MERMAID_KROKI_BACKGROUND or "").strip()
        if bg:
            # Insert a background rect as the first child to avoid transparent SVGs
            # becoming unreadable in Word dark mode.
            msvg2 = re.search(r"<svg\b[^>]*>", s2, re.IGNORECASE)
            if msvg2 and "data-owui-bg" not in s2:
                tag_now = msvg2.group(0)
                vb = re.search(
                    r'viewBox="([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)\s+'
                    r'([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)\s+'
                    r'([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)\s+'
                    r'([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)"',
                    tag_now,
                )
                rect = None
                if vb:
                    rect = (
                        f'<rect data-owui-bg="1" x="{vb.group(1)}" y="{vb.group(2)}" '
                        f'width="{vb.group(3)}" height="{vb.group(4)}" fill="{bg}"/>' 
                    )
                else:
                    rect = f'<rect data-owui-bg="1" x="0" y="0" width="100%" height="100%" fill="{bg}"/>'
                insert_at = msvg2.end()
                s2 = s2[:insert_at] + rect + s2[insert_at:]

        return s2.encode("utf-8", errors="replace")

        # unreachable

    def _svg_aspect_ratio(self, svg_bytes: bytes) -> Optional[float]:
        try:
            s = svg_bytes.decode("utf-8", errors="strict")
        except Exception:
            return None

        m = re.search(r"<svg\b[^>]*>", s, re.IGNORECASE)
        if not m:
            return None
        tag = m.group(0)

        def _num(attr: str) -> Optional[float]:
            mm = re.search(rf'{attr}="([^"]+)"', tag, re.IGNORECASE)
            if not mm:
                return None
            v = mm.group(1).strip()
            # Ignore percentages; they don't define intrinsic ratio reliably.
            if v.endswith("%"):
                return None
            # Extract leading number (px, pt, etc. are ignored for ratio).
            mn = re.match(r"^([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)", v)
            if not mn:
                return None
            try:
                return float(mn.group(1))
            except Exception:
                return None

        w = _num("width")
        h = _num("height")
        if w and h and w > 0 and h > 0:
            return w / h

        mvb = re.search(
            r'viewBox="([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)\s+'
            r'([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)\s+'
            r'([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)\s+'
            r'([+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?)"',
            tag,
        )
        if not mvb:
            return None
        try:
            vbw = float(mvb.group(3))
            vbh = float(mvb.group(4))
        except Exception:
            return None
        if vbw <= 0 or vbh <= 0:
            return None
        return vbw / vbh

    def _make_transparent_png(self, width_px: int, height_px: int) -> bytes:
        # Minimal RGBA PNG with transparent pixels.
        w = max(1, int(width_px))
        h = max(1, int(height_px))

        def _chunk(typ: bytes, data: bytes) -> bytes:
            return (
                struct.pack(">I", len(data))
                + typ
                + data
                + struct.pack(">I", zlib.crc32(typ + data) & 0xFFFFFFFF)
            )

        # Each scanline: filter byte 0 + RGBA bytes (all zero).
        row = b"\x00" + (b"\x00" * (w * 4))
        raw = row * h
        compressed = zlib.compress(raw, level=9)

        signature = b"\x89PNG\r\n\x1a\n"
        ihdr = struct.pack(">IIBBBBB", w, h, 8, 6, 0, 0, 0)  # 8-bit RGBA
        return signature + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", compressed) + _chunk(b"IEND", b"")

    def _transparent_png_for_svg(self, svg_bytes: bytes) -> bytes:
        ratio = self._svg_aspect_ratio(svg_bytes)
        if not ratio or ratio <= 0:
            return _TRANSPARENT_1PX_PNG

        max_dim = 800
        if ratio >= 1.0:
            w = max_dim
            h = max(1, int(round(max_dim / ratio)))
        else:
            h = max_dim
            w = max(1, int(round(max_dim * ratio)))

        return self._make_transparent_png(w, h)

    async def _render_mermaid_to_svg(
        self, client: httpx.AsyncClient, source: str
    ) -> _MermaidOutcome:
        mermaid_text = self._prepare_mermaid_for_kroki(source)
        data = mermaid_text.encode("utf-8", errors="replace")
        if len(data) > self.valves.MERMAID_KROKI_MAX_REQUEST_BYTES:
            return _MermaidOutcome(
                kind="code",
                error_classification="renderer_request_too_large",
                error_detail=f"{len(data)} bytes > {self.valves.MERMAID_KROKI_MAX_REQUEST_BYTES}",
            )

        base = (self.valves.MERMAID_KROKI_BASE_URL or "").strip().rstrip("/")
        url = f"{base}/mermaid/svg"
        timeout = httpx.Timeout(self.valves.MERMAID_KROKI_TIMEOUT_S)

        try:
            async with client.stream(
                "POST",
                url,
                content=data,
                headers={"Content-Type": "text/plain", "Accept": "image/svg+xml"},
                timeout=timeout,
            ) as resp:
                status = resp.status_code
                content_type = (resp.headers.get("content-type") or "").lower()

                if status != 200:
                    snippet = await self._read_text_snippet(resp, limit=300)
                    classification = "renderer_http_error"
                    if status in (400, 422):
                        classification = "renderer_http_syntax_error"
                    elif status == 429:
                        classification = "renderer_http_rate_limited"
                    elif 500 <= status <= 599:
                        classification = "renderer_http_server_error"

                    logger.warning(
                        f"Mermaid renderer non-200: status={status} content-type={content_type!r} body={snippet!r}"
                    )
                    return _MermaidOutcome(
                        kind="code",
                        error_classification=classification,
                        error_detail=str(status),
                    )

                svg = await self._read_limited_bytes(
                    resp, limit=self.valves.MERMAID_KROKI_MAX_RESPONSE_BYTES
                )
                if not self._looks_like_svg(svg, content_type):
                    snippet = svg[:300].decode("utf-8", errors="replace")
                    logger.warning(
                        f"Mermaid renderer returned unexpected payload: content-type={content_type!r} snippet={snippet!r}"
                    )
                    return _MermaidOutcome(
                        kind="code",
                        error_classification="renderer_unexpected_payload",
                        error_detail=content_type or "unknown",
                    )

                svg = self._pad_svg_viewbox(svg)
                svg = self._normalize_svg_for_word(svg)
                return _MermaidOutcome(kind="image", svg_bytes=svg)

        except httpx.TimeoutException:
            return _MermaidOutcome(kind="code", error_classification="renderer_timeout")
        except httpx.RequestError as exc:
            logger.warning(f"Mermaid renderer network error: {exc}")
            return _MermaidOutcome(kind="code", error_classification="renderer_network_error")
        except _MermaidResponseTooLarge:
            return _MermaidOutcome(
                kind="code", error_classification="renderer_response_too_large"
            )
        except Exception as exc:
            logger.exception(f"Mermaid renderer unexpected error: {exc}")
            return _MermaidOutcome(kind="code", error_classification="renderer_network_error")

    async def _read_limited_bytes(
        self, resp: httpx.Response, limit: int
    ) -> bytes:
        buf = bytearray()
        async for chunk in resp.aiter_bytes():
            if not chunk:
                continue
            buf.extend(chunk)
            if len(buf) > limit:
                raise _MermaidResponseTooLarge()
        return bytes(buf)

    async def _read_text_snippet(self, resp: httpx.Response, limit: int = 300) -> str:
        try:
            data = await self._read_limited_bytes(resp, limit=min(4096, max(1, limit * 8)))
        except Exception:
            return ""
        return data[:limit].decode("utf-8", errors="replace")

    def _looks_like_png(self, data: bytes, content_type: str) -> bool:
        if content_type and "image/png" in content_type:
            return True
        return data.startswith(b"\x89PNG\r\n\x1a\n")

    def _insert_mermaid_outcome(
        self,
        doc: Document,
        mermaid_source: str,
        outcome: _MermaidOutcome,
        info_raw: str,
        attrs: List[str],
    ):
        caption_title: Optional[str] = (
            self._extract_mermaid_title(mermaid_source)
            if self.valves.MERMAID_CAPTIONS_ENABLE
            else None
        )

        if outcome.kind == "image" and outcome.png_bytes:
            try:
                shape = self._insert_png_image(doc, outcome.png_bytes)
                if outcome.svg_bytes:
                    self._attach_svg_blip(doc, shape, outcome.svg_bytes)
                self._add_mermaid_caption(doc, caption_title)
                return
            except Exception as exc:
                logger.warning(f"Image embedding failed; falling back to code: {exc}")
                outcome = _MermaidOutcome(
                    kind="code",
                    error_classification="image_embedding_error",
                    error_detail=str(exc),
                )

        # Code fallback
        self.add_code_block(doc, mermaid_source, "mermaid")

    def _extract_mermaid_title(self, source: str) -> Optional[str]:
        lines = self._normalize_mermaid_text(source).split("\n")
        header_found = False
        for raw in lines:
            line = raw.strip()
            if not line:
                continue
            if line.startswith("%%{") and line.endswith("}%%"):
                continue
            if line.startswith("%%"):
                continue
            # diagram header line
            if not header_found:
                header_found = True
                # pie can embed title on the header line
                if line.lower().startswith("pie"):
                    m = re.match(r"^pie\s+(.+)$", line, re.IGNORECASE)
                    if m:
                        tail = m.group(1)
                        mt = re.search(r"\btitle\s+(.+)$", tail, re.IGNORECASE)
                        if mt:
                            return mt.group(1).strip().strip('"').strip("'")
                continue

            # title "Foo" / title Foo
            m = re.match(r'^title\s+"(.+)"\s*$', line, re.IGNORECASE)
            if m:
                return m.group(1).strip()
            m = re.match(r"^title\s+(.+)$", line, re.IGNORECASE)
            if m:
                return m.group(1).strip().strip('"').strip("'")
        return None

    def _strip_mermaid_title_for_render(self, source: str) -> str:
        """
        Removes Mermaid title directives from the source before sending to renderers.
        Captions already carry the title, and Kroki embedding can be undesirable.
        """
        lines = self._normalize_mermaid_text(source).split("\n")
        out: List[str] = []
        header_found = False
        title_stripped = False
        meaningful_after_header = False

        for raw in lines:
            line = raw.rstrip("\n")
            stripped = line.strip()

            if not stripped:
                out.append(line)
                continue

            if stripped.startswith("%%{") and stripped.endswith("}%%"):
                out.append(line)
                continue
            if stripped.startswith("%%"):
                out.append(line)
                continue

            if not header_found:
                header_found = True
                # pie title "..." is on header line; strip from 'title' onward.
                if stripped.lower().startswith("pie") and re.search(r"\btitle\b", stripped, re.IGNORECASE):
                    # Keep everything up to the 'title' keyword.
                    parts = re.split(r"\btitle\b", stripped, maxsplit=1, flags=re.IGNORECASE)
                    cleaned = (parts[0] if parts else stripped).strip()
                    out.append(cleaned if cleaned else "pie")
                    title_stripped = True
                    continue
                out.append(line)
                continue

            if not title_stripped and not meaningful_after_header:
                # Strip a standalone title directive line early in the diagram.
                if re.match(r'^title\s+(".+"|.+)$', stripped, re.IGNORECASE):
                    title_stripped = True
                    continue

            # Consider this a meaningful content line after header.
            meaningful_after_header = True
            out.append(line)

        return "\n".join(out).strip() + "\n"

    def _ensure_caption_style(self, doc: Document) -> str:
        if self._caption_style_name is not None:
            return self._caption_style_name

        style_name = (self.valves.MERMAID_CAPTION_STYLE or "").strip()
        if style_name == "":
            # Empty means: do not apply a caption style.
            self._caption_style_name = ""
            return ""

        # Prefer existing style if present.
        try:
            _ = doc.styles[style_name]
            self._caption_style_name = style_name
            return style_name
        except KeyError:
            pass

        # If user requested "Caption" but it's missing, create a safe custom style name.
        if style_name.lower() == "caption":
            style_name = "OWUI Caption"

        try:
            _ = doc.styles[style_name]
            self._caption_style_name = style_name
            return style_name
        except KeyError:
            pass

        try:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = "Calibri"
            style.font.size = Pt(9)
            style.font.color.rgb = RGBColor(80, 80, 80)
            style.paragraph_format.space_before = Pt(2)
            style.paragraph_format.space_after = Pt(8)
            self._caption_style_name = style_name
            return style_name
        except Exception:
            self._caption_style_name = "Normal"
            return "Normal"

    def _add_mermaid_caption(self, doc: Document, title: Optional[str]):
        if not self.valves.MERMAID_CAPTIONS_ENABLE:
            return

        prefix = (self.valves.MERMAID_CAPTION_PREFIX or "").strip()
        if prefix == "" and not title:
            return

        self._mermaid_figure_counter += 1
        if prefix == "":
            caption = title or ""
        else:
            base = f"{prefix} {self._mermaid_figure_counter}"
            caption = f"{base}: {title}" if title else base
        if caption == "":
            return

        para = doc.add_paragraph()
        style_name = self._ensure_caption_style(doc)
        if style_name:
            para.style = style_name
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_formatted_text(para, caption)

    def _available_block_width(self, doc: Document):
        section = doc.sections[0]
        return section.page_width - section.left_margin - section.right_margin

    def _insert_png_image(self, doc: Document, png_bytes: bytes):
        width = self._available_block_width(doc)
        return doc.add_picture(cast(Any, io.BytesIO(png_bytes)), width=width)

    def _attach_svg_blip(self, doc: Document, inline_shape: Any, svg_bytes: bytes):
        if not svg_bytes:
            return

        try:
            pkg = doc.part.package
            partname = pkg.next_partname("/word/media/image%d.svg")
            from docx.opc.part import Part

            svg_part = Part(partname, "image/svg+xml", svg_bytes)
            rid_svg = doc.part.relate_to(svg_part, RT.IMAGE)

            inline = inline_shape._inline
            blips = inline.xpath(".//a:blip")
            if not blips:
                return
            blip = blips[0]

            existing = blip.xpath(".//asvg:svgBlip")
            if existing:
                existing[0].set(qn("r:embed"), rid_svg)
                return

            extLst = OxmlElement("a:extLst")
            ext = OxmlElement("a:ext")
            ext.set("uri", "{96DAC541-7B7A-43D3-8B79-37D633B846F1}")

            svgBlip = OxmlElement("asvg:svgBlip")
            svgBlip.set(qn("r:embed"), rid_svg)
            ext.append(svgBlip)
            extLst.append(ext)
            blip.append(extLst)
        except Exception as exc:
            logger.warning(f"Failed to attach SVG blip; keeping PNG fallback: {exc}")

    # (Mermaid warning paragraphs removed)

    def set_document_default_font(self, doc: Document):
        """Set document default font (Word-like defaults)."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Set paragraph format
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.space_after = Pt(6)

    def add_heading(self, doc: Document, text: str, level: int):
        """Add heading"""
        # Word heading levels start from 0, Markdown from 1
        heading_level = min(level, 9)  # Word supports up to Heading 9
        heading = doc.add_heading(level=heading_level)

        # Parse and add formatted text
        self.add_formatted_text(heading, text)

    def add_paragraph(self, doc: Document, text: str):
        """Add paragraph with inline formatting support"""
        paragraph = doc.add_paragraph()
        self.add_formatted_text(paragraph, text)

    def add_formatted_text(self, paragraph, text: str):
        """
        Parse Markdown inline formatting and add to paragraph.
        Supports: bold, italic, inline code, links, strikethrough, auto-link URLs,
        and inline LaTeX math \\(...\\) when MATH_ENABLE is on.
        """
        self._add_inline_segments(paragraph, text or "", bold=False, italic=False, strike=False)

    def _add_text_run(self, paragraph, s: str, bold: bool, italic: bool, strike: bool):
        if not s:
            return
        run = paragraph.add_run(s)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if strike:
            run.font.strike = True

    def _add_inline_code(self, paragraph, s: str):
        if s == "":
            return

        def _add_code_run(chunk: str):
            if not chunk:
                return
            run = paragraph.add_run(chunk)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "E8E8E8")
            run._element.rPr.append(shading)

        i = 0
        for m in _AUTO_URL_RE.finditer(s):
            start, end = m.span()
            if start > i:
                _add_code_run(s[i:start])

            raw = m.group(0)
            trimmed = raw
            while trimmed and trimmed[-1] in ".,;:!?)]}":
                trimmed = trimmed[:-1]
            suffix = raw[len(trimmed) :]

            normalized = self._normalize_url(trimmed)
            if normalized:
                self._add_hyperlink_code(paragraph, display_text=trimmed, url=normalized)
            else:
                _add_code_run(raw)

            if suffix:
                _add_code_run(suffix)

            i = end

        if i < len(s):
            _add_code_run(s[i:])

    def _add_hyperlink_code(self, paragraph, display_text: str, url: str):
        u = self._normalize_url(url)
        if not u:
            self._add_inline_code(paragraph, display_text)
            return

        part = getattr(paragraph, "part", None)
        if part is None or not hasattr(part, "relate_to"):
            self._add_inline_code(paragraph, display_text)
            return

        r_id = part.relate_to(u, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        rFonts.set(qn("w:eastAsia"), "Consolas")
        rPr.append(rFonts)

        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "20")  # 10pt
        rPr.append(sz)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), "20")
        rPr.append(sz_cs)

        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E8E8E8")
        rPr.append(shading)

        new_run.append(rPr)

        t = OxmlElement("w:t")
        t.text = display_text
        new_run.append(t)

        hyperlink.append(new_run)
        cast(Any, paragraph)._p.append(hyperlink)

    def _add_inline_segments(self, paragraph, text: str, bold: bool, italic: bool, strike: bool):
        i = 0
        n = len(text)

        def next_special(start: int) -> int:
            candidates = []
            for ch in ("`", "[", "*", "_", "~"):
                idx = text.find(ch, start)
                if idx != -1:
                    candidates.append(idx)
            idx = text.find(r"\(", start)
            if idx != -1:
                candidates.append(idx)
            idx = text.find("http://", start)
            if idx != -1:
                candidates.append(idx)
            idx = text.find("https://", start)
            if idx != -1:
                candidates.append(idx)
            idx = text.find("www.", start)
            if idx != -1:
                candidates.append(idx)
            return min(candidates) if candidates else n

        while i < n:
            if text[i] == "`":
                j = text.find("`", i + 1)
                if j != -1:
                    self._add_inline_code(paragraph, text[i + 1 : j])
                    i = j + 1
                    continue

            if text.startswith(r"\(", i):
                j = text.find(r"\)", i + 2)
                if j != -1:
                    self._add_inline_equation(paragraph, text[i + 2 : j], bold=bold, italic=italic, strike=strike)
                    i = j + 2
                    continue

            if text.startswith("~~", i):
                j = text.find("~~", i + 2)
                if j != -1:
                    self._add_inline_segments(paragraph, text[i + 2 : j], bold=bold, italic=italic, strike=True)
                    i = j + 2
                    continue

            if text.startswith("**", i):
                j = text.find("**", i + 2)
                if j != -1:
                    self._add_inline_segments(paragraph, text[i + 2 : j], bold=True, italic=italic, strike=strike)
                    i = j + 2
                    continue

            if text.startswith("__", i):
                j = text.find("__", i + 2)
                if j != -1:
                    self._add_inline_segments(paragraph, text[i + 2 : j], bold=True, italic=italic, strike=strike)
                    i = j + 2
                    continue

            if text[i] == "*" and (i + 1 >= n or text[i + 1] != "*"):
                j = text.find("*", i + 1)
                if j != -1:
                    self._add_inline_segments(paragraph, text[i + 1 : j], bold=bold, italic=True, strike=strike)
                    i = j + 1
                    continue

            if text[i] == "_" and (i + 1 >= n or text[i + 1] != "_"):
                j = text.find("_", i + 1)
                if j != -1:
                    self._add_inline_segments(paragraph, text[i + 1 : j], bold=bold, italic=True, strike=strike)
                    i = j + 1
                    continue

            if text[i] == "[":
                close = text.find("]", i + 1)
                if close != -1 and close + 1 < n and text[close + 1] == "(":
                    close_paren = text.find(")", close + 2)
                    if close_paren != -1:
                        label = text[i + 1 : close]
                        url = text[close + 2 : close_paren]
                        self._add_hyperlink(paragraph, label, url)
                        i = close_paren + 1
                        continue
                # Citation marker like [12] -> internal link to References.
                if close != -1:
                    inner = text[i + 1 : close].strip()
                    if inner.isdigit():
                        idx = int(inner)
                        anchor = self._citation_anchor_by_index.get(idx)
                        if anchor:
                            self._add_internal_hyperlink(paragraph, f"[{idx}]", anchor)
                            i = close + 1
                            continue

            m = _AUTO_URL_RE.match(text, i)
            if m:
                raw = m.group(0)
                trimmed = raw
                while trimmed and trimmed[-1] in ".,;:!?)]}":
                    trimmed = trimmed[:-1]
                suffix = raw[len(trimmed) :]

                normalized = self._normalize_url(trimmed)
                if normalized:
                    # Display the original (trimmed) text; use normalized URL as the target.
                    self._add_hyperlink(paragraph, trimmed, normalized, display_text=trimmed)
                else:
                    self._add_text_run(paragraph, raw, bold, italic, strike)
                    i += len(raw)
                    continue

                if suffix:
                    self._add_text_run(paragraph, suffix, bold, italic, strike)
                i += len(raw)
                continue

            j = next_special(i)
            if j == i:
                # Unmatched special character; treat literally to avoid infinite loops.
                self._add_text_run(paragraph, text[i], bold, italic, strike)
                i += 1
            else:
                self._add_text_run(paragraph, text[i:j], bold, italic, strike)
                i = j

    def _normalize_url(self, url: str) -> str:
        u = (url or "").strip()
        if u.lower().startswith("www."):
            u = "https://" + u

        # Trim common trailing punctuation that often follows URLs in prose.
        while u and u[-1] in ".,;:!?)]}":
            u = u[:-1]
        return u

    def _add_hyperlink(self, paragraph, text: str, url: str, display_text: Optional[str] = None):
        u = self._normalize_url(url)
        if not u:
            paragraph.add_run(display_text or text)
            return

        part = getattr(paragraph, "part", None)
        if part is None or not hasattr(part, "relate_to"):
            # Fallback if relationship API isn't available.
            run = paragraph.add_run(display_text or text)
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.underline = True
            return

        r_id = part.relate_to(u, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(color)

        u_el = OxmlElement("w:u")
        u_el.set(qn("w:val"), "single")
        rPr.append(u_el)

        new_run.append(rPr)

        t = OxmlElement("w:t")
        t.text = display_text or text
        new_run.append(t)

        hyperlink.append(new_run)
        cast(Any, paragraph)._p.append(hyperlink)

    def _add_inline_equation(self, paragraph, latex: str, bold: bool = False, italic: bool = False, strike: bool = False):
        latex = (latex or "").strip()
        if not latex:
            return

        if not self.valves.MATH_ENABLE or not LATEX_MATH_AVAILABLE:
            self._add_text_run(paragraph, f"\\({latex}\\)", bold=bold, italic=italic, strike=strike)
            return

        try:
            mathml = latex_to_mathml(latex)
            omml = mathml2omml.convert(mathml)
            o_math = self._omml_oMath_element(omml)
            run = paragraph.add_run()
            run.bold = bold
            run.italic = italic
            run.font.strike = strike
            cast(Any, run)._r.append(o_math)
        except Exception as exc:
            logger.warning(f"Inline math conversion failed; keeping literal: {exc}")
            self._add_text_run(paragraph, f"\\({latex}\\)", bold=bold, italic=italic, strike=strike)

    def _omml_oMath_element(self, omml: str):
        # Ensure the OMML element declares the math namespace so parse_xml works.
        m_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        s = (omml or "").strip()
        if s.startswith("<m:oMath>") and s.endswith("</m:oMath>"):
            inner = s[len("<m:oMath>") : -len("</m:oMath>")]
            s = f'<m:oMath xmlns:m="{m_ns}">{inner}</m:oMath>'
        elif s.startswith("<m:oMath") and "xmlns:m=" not in s.split(">", 1)[0]:
            s = s.replace("<m:oMath", f'<m:oMath xmlns:m="{m_ns}"', 1)
        return parse_xml(s)

    def add_code_block(self, doc: Document, code: str, language: str = ""):
        """Add code block with syntax highlighting"""
        # Token color mapping (based on common IDE themes)
        TOKEN_COLORS = {
            Token.Keyword: RGBColor(0, 92, 197),  # macOS blue - keywords
            Token.Keyword.Constant: RGBColor(0, 92, 197),
            Token.Keyword.Declaration: RGBColor(0, 92, 197),
            Token.Keyword.Namespace: RGBColor(0, 92, 197),
            Token.Keyword.Type: RGBColor(0, 92, 197),
            Token.Name.Function: RGBColor(0, 0, 0),  # Functions stay black
            Token.Name.Class: RGBColor(38, 82, 120),  # Deep cyan-blue - classes
            Token.Name.Decorator: RGBColor(170, 51, 0),  # Warm orange - decorators
            Token.Name.Builtin: RGBColor(0, 110, 71),  # Deep green - builtins
            Token.String: RGBColor(196, 26, 22),  # Red - strings
            Token.String.Doc: RGBColor(109, 120, 133),  # Gray - docstrings
            Token.Comment: RGBColor(109, 120, 133),  # Gray - comments
            Token.Comment.Single: RGBColor(109, 120, 133),
            Token.Comment.Multiline: RGBColor(109, 120, 133),
            Token.Number: RGBColor(28, 0, 207),  # Indigo - numbers
            Token.Number.Integer: RGBColor(28, 0, 207),
            Token.Number.Float: RGBColor(28, 0, 207),
            Token.Operator: RGBColor(90, 99, 120),  # Gray-blue - operators
            Token.Punctuation: RGBColor(0, 0, 0),  # Black - punctuation
        }

        def get_token_color(token_type):
            """Recursively find token color"""
            while token_type:
                if token_type in TOKEN_COLORS:
                    return TOKEN_COLORS[token_type]
                token_type = token_type.parent
            return None

        # Add language label if available
        if language:
            lang_para = doc.add_paragraph()
            lang_para.paragraph_format.space_before = Pt(6)
            lang_para.paragraph_format.space_after = Pt(0)
            lang_para.paragraph_format.left_indent = Cm(0.5)
            lang_run = lang_para.add_run(language.upper())
            lang_run.font.name = "Consolas"
            lang_run.font.size = Pt(8)
            lang_run.font.color.rgb = RGBColor(100, 100, 100)
            lang_run.font.bold = True

        # Add code block paragraph
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.5)
        paragraph.paragraph_format.space_before = Pt(3) if language else Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

        # Add light gray background
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F7F7F7")
        paragraph._element.pPr.append(shading)

        # Try to use Pygments for syntax highlighting
        if PYGMENTS_AVAILABLE and language:
            try:
                lexer = get_lexer_by_name(language, stripall=False)
            except Exception:
                lexer = TextLexer()

            tokens = list(lex(code, lexer))

            for token_type, token_value in tokens:
                if not token_value:
                    continue
                run = paragraph.add_run(token_value)
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                run.font.size = Pt(10)

                # Apply color
                color = get_token_color(token_type)
                if color:
                    run.font.color.rgb = color

                # Bold keywords
                if token_type in Token.Keyword:
                    run.font.bold = True
        else:
            # No syntax highlighting, plain text display
            run = paragraph.add_run(code)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10)

    def add_table(self, doc: Document, table_lines: List[str]):
        """Add Markdown table with sane Word sizing/spacing, alignment, and hyperlinks/math support in cells."""
        if len(table_lines) < 2:
            return

        header_fill = "F2F2F2"
        zebra_fill = "FBFBFB"

        def _split_row(line: str) -> List[str]:
            # Keep empty cells, trim surrounding pipes.
            raw = line.strip().strip("|")
            return [c.strip() for c in raw.split("|")]

        def _is_separator_row(cells: List[str]) -> bool:
            # Markdown separator: --- / :--- / ---: / :---:
            if not cells:
                return False
            ok = 0
            for c in cells:
                c = c.strip()
                if re.fullmatch(r":?-{3,}:?", c):
                    ok += 1
            return ok == len(cells)

        def _col_align(cell: str) -> WD_ALIGN_PARAGRAPH:
            s = (cell or "").strip()
            if s.startswith(":") and s.endswith(":"):
                return WD_ALIGN_PARAGRAPH.CENTER
            if s.endswith(":"):
                return WD_ALIGN_PARAGRAPH.RIGHT
            return WD_ALIGN_PARAGRAPH.LEFT

        def _set_cell_shading(cell, fill: str):
            tc_pr = cell._element.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), fill)
            tc_pr.append(shd)

        raw_rows = [_split_row(l) for l in table_lines if l.strip().startswith("|")]
        if not raw_rows:
            return

        sep_idx = 1 if len(raw_rows) > 1 and _is_separator_row(raw_rows[1]) else -1
        header = raw_rows[0]
        body = raw_rows[sep_idx + 1 :] if sep_idx >= 0 else raw_rows[1:]

        num_cols = max(len(header), *(len(r) for r in body)) if body else len(header)
        header = header + [""] * (num_cols - len(header))
        body = [r + [""] * (num_cols - len(r)) for r in body]

        aligns = [_col_align(c) for c in (raw_rows[1] if sep_idx == 1 else [""] * num_cols)]

        table = doc.add_table(rows=1 + len(body), cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cast(Any, table).autofit = False

        # Cell margins (twips): smaller padding for compact tables.
        self._set_table_cell_margins(table, top=60, bottom=60, left=90, right=90)

        # Column widths: proportional to content, bounded, then normalized to page width.
        available_width = int(self._available_block_width(doc))
        min_col = max(int(Inches(0.55)), available_width // max(1, num_cols * 3))

        def _plain_len(s: str) -> int:
            t = re.sub(r"`([^`]+)`", r"\\1", s or "")
            t = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\\1", t)
            t = re.sub(r"\\s+", " ", t).strip()
            return len(t)

        weights: List[int] = []
        for ci in range(num_cols):
            max_len = _plain_len(header[ci])
            for r in body:
                max_len = max(max_len, _plain_len(r[ci]))
            weights.append(max(1, min(max_len, 40)))

        sum_w = sum(weights) or 1
        widths = [max(min_col, int(available_width * w / sum_w)) for w in weights]
        total = sum(widths)
        if total > available_width:
            even = max(1, available_width // max(1, num_cols))
            widths = [even] * num_cols
            total = sum(widths)
        if total < available_width:
            rem = available_width - total
            order = sorted(range(num_cols), key=lambda i: weights[i], reverse=True)
            oi = 0
            while rem > 0 and order:
                widths[order[oi % len(order)]] += 1
                rem -= 1
                oi += 1

        for ci, w in enumerate(widths):
            table.columns[ci].width = w
            for row in table.rows:
                row.cells[ci].width = w

        def _format_cell_paragraph(para, align: WD_ALIGN_PARAGRAPH):
            para.alignment = align
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

        def _fill_cell(cell, text: str, align: WD_ALIGN_PARAGRAPH, bold: bool = False):
            cell.text = ""
            parts = [p for p in re.split(r"(?:<br\\s*/?>|\\n)", text or "") if p is not None]
            if not parts:
                parts = [""]
            for pi, part in enumerate(parts):
                para = cell.paragraphs[0] if pi == 0 else cell.add_paragraph()
                _format_cell_paragraph(para, align)
                self.add_formatted_text(para, part)
                for run in para.runs:
                    run.font.size = Pt(9)
                    if bold:
                        run.bold = True

        # Header row
        header_row = table.rows[0]
        self._set_table_header_row_repeat(header_row)
        for ci in range(num_cols):
            cell = header_row.cells[ci]
            _set_cell_shading(cell, header_fill)
            _fill_cell(cell, header[ci], aligns[ci] if ci < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT, bold=True)

        # Body rows
        for ri, row_data in enumerate(body, start=1):
            row = table.rows[ri]
            for ci in range(num_cols):
                cell = row.cells[ci]
                if (ri % 2) == 0:
                    _set_cell_shading(cell, zebra_fill)
                _fill_cell(cell, row_data[ci], aligns[ci] if ci < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT)

    def _set_table_cell_margins(self, table, top: int, bottom: int, left: int, right: int):
        tbl_pr = cast(Any, table)._tbl.tblPr
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        for tag, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
            el = OxmlElement(f"w:{tag}")
            el.set(qn("w:w"), str(int(val)))
            el.set(qn("w:type"), "dxa")
            tbl_cell_mar.append(el)
        tbl_pr.append(tbl_cell_mar)

    def _set_table_header_row_repeat(self, row):
        tr_pr = row._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)

    def add_list_to_doc(
        self, doc: Document, items: List[Tuple[int, str]], list_type: str
    ):
        """Add list"""
        for indent, text in items:
            paragraph = doc.add_paragraph()

            if list_type == "unordered":
                # Unordered list with bullets
                paragraph.style = "List Bullet"
            else:
                # Ordered list with numbers
                paragraph.style = "List Number"

            # Set indent
            paragraph.paragraph_format.left_indent = Cm(0.5 * (indent + 1))

            # Add formatted text
            self.add_formatted_text(paragraph, text)

    def add_horizontal_rule(self, doc: Document):
        """Add horizontal rule"""
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)

        # Add bottom border as horizontal rule
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_blockquote(self, doc: Document, text: str):
        """Add blockquote with left border and gray background"""
        for line in text.split("\n"):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(1.0)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)

            # Add left border
            pPr = paragraph._element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "24")  # Border thickness
            left.set(qn("w:space"), "4")  # Space between border and text
            left.set(qn("w:color"), "CCCCCC")  # Gray border
            pBdr.append(left)
            pPr.append(pBdr)

            # Add light gray background
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F9F9F9")
            pPr.append(shading)

            # Add formatted text
            self.add_formatted_text(paragraph, line)

            # Set font to italic gray
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(85, 85, 85)  # Dark gray text
                run.italic = True
