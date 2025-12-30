"""
title: 导出为 Word
author: Fu-Jie
author_url: https://github.com/Fu-Jie
funding_url: https://github.com/Fu-Jie/awesome-openwebui
version: 0.1.0
icon_url: data:image/svg+xml;base64,PHN2ZwogIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyIKICB3aWR0aD0iMjQiCiAgaGVpZ2h0PSIyNCIKICB2aWV3Qm94PSIwIDAgMjQgMjQiCiAgZmlsbD0ibm9uZSIKICBzdHJva2U9ImN1cnJlbnRDb2xvciIKICBzdHJva2Utd2lkdGg9IjIiCiAgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIgogIHN0cm9rZS1saW5lam9pbj0icm91bmQiCj4KICA8cGF0aCBkPSJNNiAyMmEyIDIgMCAwIDEtMi0yVjRhMiAyIDAgMCAxIDItMmg4YTIuNCAyLjQgMCAwIDEgMS43MDQuNzA2bDMuNTg4IDMuNTg4QTIuNCAyLjQgMCAwIDEgMjAgOHYxMmEyIDIgMCAwIDEtMiAyeiIgLz4KICA8cGF0aCBkPSJNMTQgMnY1YTEgMSAwIDAgMCAxIDFoNSIgLz4KICA8cGF0aCBkPSJNMTAgOUg4IiAvPgogIDxwYXRoIGQ9Ik0xNiAxM0g4IiAvPgogIDxwYXRoIGQ9Ik0xNiAxN0g4IiAvPgo8L3N2Zz4K
requirements: python-docx==1.1.2, Pygments>=2.15.0
description: 将当前对话内容从 Markdown 转换并导出为 Word (.docx) 文件，支持代码语法高亮和引用块。
"""

import os
import re
import base64
import datetime
import io
import asyncio
import logging
from typing import Optional, Callable, Awaitable, Any, List, Tuple
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from open_webui.models.chats import Chats

# Pygments for syntax highlighting
try:
    from pygments import lex
    from pygments.lexers import get_lexer_by_name, TextLexer
    from pygments.token import Token

    PYGMENTS_AVAILABLE = True
except ImportError:
    PYGMENTS_AVAILABLE = False


logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)


class Action:
    def __init__(self):
        pass

    async def _send_notification(self, emitter: Callable, type: str, content: str):
        await emitter(
            {"type": "notification", "data": {"type": type, "content": content}}
        )

    async def action(
        self,
        body: dict,
        __user__=None,
        __event_emitter__=None,
        __event_call__: Optional[Callable[[Any], Awaitable[None]]] = None,
        __metadata__: Optional[dict] = None,
    ):
        logger.info(f"action:{__name__}")

        # 解析用户信息
        if isinstance(__user__, (list, tuple)):
            user_language = (
                __user__[0].get("language", "zh-CN") if __user__ else "zh-CN"
            )
            user_name = __user__[0].get("name", "用户") if __user__[0] else "用户"
            user_id = (
                __user__[0]["id"]
                if __user__ and "id" in __user__[0]
                else "unknown_user"
            )
        elif isinstance(__user__, dict):
            user_language = __user__.get("language", "zh-CN")
            user_name = __user__.get("name", "用户")
            user_id = __user__.get("id", "unknown_user")

        if __event_emitter__:
            last_assistant_message = body["messages"][-1]

            await __event_emitter__(
                {
                    "type": "status",
                    "data": {"description": "正在转换为 Word 文档...", "done": False},
                }
            )

            try:
                message_content = last_assistant_message["content"]

                if not message_content or not message_content.strip():
                    await self._send_notification(
                        __event_emitter__, "error", "没有找到可导出的内容！"
                    )
                    return

                # 生成文件名（优先对话标题；若缺失则通过 chat_id 查询；再到 Markdown 标题；最后用户+日期）
                chat_id = self.extract_chat_id(body, __metadata__)
                chat_title = self.extract_chat_title(body)
                if not chat_title and chat_id:
                    chat_title = await self.fetch_chat_title(chat_id, user_id)
                title = self.extract_title(message_content)
                current_datetime = datetime.datetime.now()
                formatted_date = current_datetime.strftime("%Y%m%d")

                if chat_title:
                    filename = f"{self.clean_filename(chat_title)}.docx"
                elif title:
                    filename = f"{self.clean_filename(title)}.docx"
                else:
                    filename = f"{user_name}_{formatted_date}.docx"

                # 创建 Word 文档；若正文无一级标题，使用对话标题作为一级标题
                has_h1 = bool(re.search(r"^#\s+.+$", message_content, re.MULTILINE))
                doc = self.markdown_to_docx(
                    message_content, top_heading=chat_title, has_h1=has_h1
                )

                # 保存到内存
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                file_content = doc_buffer.read()
                base64_blob = base64.b64encode(file_content).decode("utf-8")

                # 触发文件下载
                if __event_call__:
                    await __event_call__(
                        {
                            "type": "execute",
                            "data": {
                                "code": f"""
                                try {{
                                    const base64Data = "{base64_blob}";
                                    const binaryData = atob(base64Data);
                                    const arrayBuffer = new Uint8Array(binaryData.length);
                                    for (let i = 0; i < binaryData.length; i++) {{
                                        arrayBuffer[i] = binaryData.charCodeAt(i);
                                    }}
                                    const blob = new Blob([arrayBuffer], {{ type: "application/vnd.openxmlformats-officedocument.wordprocessingml.document" }});
                                    const filename = "{filename}";

                                    const url = URL.createObjectURL(blob);
                                    const a = document.createElement("a");
                                    a.style.display = "none";
                                    a.href = url;
                                    a.download = filename;
                                    document.body.appendChild(a);
                                    a.click();
                                    URL.revokeObjectURL(url);
                                    document.body.removeChild(a);
                                }} catch (error) {{
                                    console.error('触发下载时出错:', error);
                                }}
                                """
                            },
                        }
                    )

                await __event_emitter__(
                    {
                        "type": "status",
                        "data": {"description": "Word 文档已导出", "done": True},
                    }
                )

                await self._send_notification(
                    __event_emitter__, "success", f"已成功导出为 {filename}"
                )

                return {"message": "下载事件已触发"}

            except Exception as e:
                print(f"Error exporting to Word: {str(e)}")
                await __event_emitter__(
                    {
                        "type": "status",
                        "data": {
                            "description": f"导出失败: {str(e)}",
                            "done": True,
                        },
                    }
                )
                await self._send_notification(
                    __event_emitter__, "error", f"导出 Word 文档时出错: {str(e)}"
                )

    def extract_title(self, content: str) -> str:
        """从 Markdown 内容提取一级/二级标题"""
        lines = content.split("\n")
        for line in lines:
            # 仅匹配 h1-h2 标题
            match = re.match(r"^#{1,2}\s+(.+)$", line.strip())
            if match:
                return match.group(1).strip()
        return ""

    def extract_chat_title(self, body: dict) -> str:
        """从请求体中提取会话标题"""
        if not isinstance(body, dict):
            return ""

        candidates = []

        for key in ("chat", "conversation"):
            if isinstance(body.get(key), dict):
                candidates.append(body.get(key, {}).get("title", ""))

        for key in ("title", "chat_title"):
            value = body.get(key)
            if isinstance(value, str):
                candidates.append(value)

        for candidate in candidates:
            if candidate and isinstance(candidate, str):
                return candidate.strip()
        return ""

    def extract_chat_id(self, body: dict, metadata: Optional[dict]) -> str:
        """从 body 或 metadata 中提取 chat_id"""
        if isinstance(body, dict):
            chat_id = body.get("chat_id") or body.get("id")
            if isinstance(chat_id, str) and chat_id.strip():
                return chat_id.strip()

            for key in ("chat", "conversation"):
                nested = body.get(key)
                if isinstance(nested, dict):
                    nested_id = nested.get("id") or nested.get("chat_id")
                    if isinstance(nested_id, str) and nested_id.strip():
                        return nested_id.strip()
        if isinstance(metadata, dict):
            chat_id = metadata.get("chat_id")
            if isinstance(chat_id, str) and chat_id.strip():
                return chat_id.strip()
        return ""

    async def fetch_chat_title(self, chat_id: str, user_id: str = "") -> str:
        """根据 chat_id 从数据库获取标题"""
        if not chat_id:
            return ""

        def _load_chat():
            if user_id:
                return Chats.get_chat_by_id_and_user_id(id=chat_id, user_id=user_id)
            return Chats.get_chat_by_id(chat_id)

        try:
            chat = await asyncio.to_thread(_load_chat)
        except Exception as exc:
            logger.warning(f"加载聊天 {chat_id} 失败: {exc}")
            return ""

        if not chat:
            return ""

        data = getattr(chat, "chat", {}) or {}
        title = data.get("title") or getattr(chat, "title", "")
        return title.strip() if isinstance(title, str) else ""

    def clean_filename(self, name: str) -> str:
        """清理文件名中的非法字符"""
        return re.sub(r'[\\/*?:"<>|]', "", name).strip()[:50]

    def markdown_to_docx(
        self, markdown_text: str, top_heading: str = "", has_h1: bool = False
    ) -> Document:
        """
        将 Markdown 文本转换为 Word 文档
        支持：标题、段落、粗体、斜体、代码块、列表、表格、链接
        """
        doc = Document()

        # 设置默认中文字体
        self.set_document_default_font(doc)

        # 若正文无一级标题且有对话标题，则作为一级标题写入
        if top_heading and not has_h1:
            self.add_heading(doc, top_heading, 1)

        lines = markdown_text.split("\n")
        i = 0
        in_code_block = False
        code_block_content = []
        code_block_lang = ""
        in_list = False
        list_items = []
        list_type = None  # 'ordered' or 'unordered'

        while i < len(lines):
            line = lines[i]

            # 处理代码块
            if line.strip().startswith("```"):
                if not in_code_block:
                    # 先处理之前积累的列表
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                        in_list = False

                    in_code_block = True
                    code_block_lang = line.strip()[3:].strip()
                    code_block_content = []
                else:
                    # 代码块结束
                    in_code_block = False
                    self.add_code_block(
                        doc, "\n".join(code_block_content), code_block_lang
                    )
                    code_block_content = []
                    code_block_lang = ""
                i += 1
                continue

            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue

            # 处理表格
            if line.strip().startswith("|") and line.strip().endswith("|"):
                # 先处理之前积累的列表
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False

                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                self.add_table(doc, table_lines)
                continue

            # 处理标题
            header_match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
            if header_match:
                # 先处理之前积累的列表
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False

                level = len(header_match.group(1))
                text = header_match.group(2)
                self.add_heading(doc, text, level)
                i += 1
                continue

            # 处理无序列表
            unordered_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
            if unordered_match:
                if not in_list or list_type != "unordered":
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                    in_list = True
                    list_type = "unordered"
                indent = len(unordered_match.group(1)) // 2
                list_items.append((indent, unordered_match.group(2)))
                i += 1
                continue

            # 处理有序列表
            ordered_match = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
            if ordered_match:
                if not in_list or list_type != "ordered":
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                    in_list = True
                    list_type = "ordered"
                indent = len(ordered_match.group(1)) // 2
                list_items.append((indent, ordered_match.group(2)))
                i += 1
                continue

            # 处理引用块
            if line.strip().startswith(">"):
                # 先处理之前积累的列表
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False

                # 收集连续的引用行
                blockquote_lines = []
                while i < len(lines) and lines[i].strip().startswith(">"):
                    # 移除开头的 > 和可能的空格
                    quote_line = re.sub(r"^>\s?", "", lines[i])
                    blockquote_lines.append(quote_line)
                    i += 1
                self.add_blockquote(doc, "\n".join(blockquote_lines))
                continue

            # 处理水平分割线
            if re.match(r"^[-*_]{3,}$", line.strip()):
                # 先处理之前积累的列表
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False

                self.add_horizontal_rule(doc)
                i += 1
                continue

            # 处理空行
            if not line.strip():
                # 列表结束
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False
                i += 1
                continue

            # 处理普通段落
            if in_list and list_items:
                self.add_list_to_doc(doc, list_items, list_type)
                list_items = []
                in_list = False

            self.add_paragraph(doc, line)
            i += 1

        # 处理剩余的列表
        if in_list and list_items:
            self.add_list_to_doc(doc, list_items, list_type)

        return doc

    def set_document_default_font(self, doc: Document):
        """设置文档默认字体，确保中英文都正常显示"""
        # 设置正文样式
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"  # 英文字体
        font.size = Pt(11)

        # 设置中文字体
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 设置段落格式
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.space_after = Pt(6)

    def add_heading(self, doc: Document, text: str, level: int):
        """添加标题"""
        # Word 标题级别从 0 开始，Markdown 从 1 开始
        heading_level = min(level, 9)  # Word 最多支持 Heading 9
        heading = doc.add_heading(level=heading_level)

        # 解析并添加格式化文本
        self.add_formatted_text(heading, text)

        # 设置中文字体
        for run in heading.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.color.rgb = RGBColor(0, 0, 0)

    def add_paragraph(self, doc: Document, text: str):
        """添加段落，支持内联格式"""
        paragraph = doc.add_paragraph()
        self.add_formatted_text(paragraph, text)

        # 设置中文字体
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def add_formatted_text(self, paragraph, text: str):
        """
        解析 Markdown 内联格式并添加到段落
        支持：粗体、斜体、行内代码、链接、删除线
        """
        # 定义格式化模式
        patterns = [
            # 粗斜体 ***text*** 或 ___text___
            (r"\*\*\*(.+?)\*\*\*|___(.+?)___", {"bold": True, "italic": True}),
            # 粗体 **text** 或 __text__
            (r"\*\*(.+?)\*\*|__(.+?)__", {"bold": True}),
            # 斜体 *text* 或 _text_
            (
                r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|(?<!_)_(?!_)(.+?)(?<!_)_(?!_)",
                {"italic": True},
            ),
            # 行内代码 `code`
            (r"`([^`]+)`", {"code": True}),
            # 链接 [text](url)
            (r"\[([^\]]+)\]\(([^)]+)\)", {"link": True}),
            # 删除线 ~~text~~
            (r"~~(.+?)~~", {"strike": True}),
        ]

        # 简化处理：逐段解析
        remaining = text
        last_end = 0

        # 合并所有匹配项
        all_matches = []

        for pattern, style in patterns:
            for match in re.finditer(pattern, text):
                # 获取匹配的文本内容
                groups = match.groups()
                matched_text = next((g for g in groups if g is not None), "")
                all_matches.append(
                    {
                        "start": match.start(),
                        "end": match.end(),
                        "text": matched_text,
                        "style": style,
                        "full_match": match.group(0),
                        "url": (
                            groups[1] if style.get("link") and len(groups) > 1 else None
                        ),
                    }
                )

        # 按位置排序
        all_matches.sort(key=lambda x: x["start"])

        # 移除重叠的匹配
        filtered_matches = []
        last_end = 0
        for m in all_matches:
            if m["start"] >= last_end:
                filtered_matches.append(m)
                last_end = m["end"]

        # 构建最终文本
        pos = 0
        for match in filtered_matches:
            # 添加匹配前的普通文本
            if match["start"] > pos:
                plain_text = text[pos : match["start"]]
                if plain_text:
                    paragraph.add_run(plain_text)

            # 添加格式化文本
            style = match["style"]
            run_text = match["text"]

            if style.get("link"):
                # 链接处理
                run = paragraph.add_run(run_text)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.font.underline = True
            elif style.get("code"):
                # 行内代码
                run = paragraph.add_run(run_text)
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
                run.font.size = Pt(10)
                # 添加背景色
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "E8E8E8")
                run._element.rPr.append(shading)
            else:
                run = paragraph.add_run(run_text)
                if style.get("bold"):
                    run.bold = True
                if style.get("italic"):
                    run.italic = True
                if style.get("strike"):
                    run.font.strike = True

            pos = match["end"]

        # 添加剩余的普通文本
        if pos < len(text):
            paragraph.add_run(text[pos:])

    def add_code_block(self, doc: Document, code: str, language: str = ""):
        """添加代码块，支持语法高亮"""
        # 语法高亮颜色映射 (基于常见的 IDE 配色)
        TOKEN_COLORS = {
            Token.Keyword: RGBColor(0, 92, 197),  # macOS 风格蓝 - 关键字
            Token.Keyword.Constant: RGBColor(0, 92, 197),
            Token.Keyword.Declaration: RGBColor(0, 92, 197),
            Token.Keyword.Namespace: RGBColor(0, 92, 197),
            Token.Keyword.Type: RGBColor(0, 92, 197),
            Token.Name.Function: RGBColor(0, 0, 0),  # 函数名保持黑色
            Token.Name.Class: RGBColor(38, 82, 120),  # 深青蓝 - 类名
            Token.Name.Decorator: RGBColor(170, 51, 0),  # 暖橙 - 装饰器
            Token.Name.Builtin: RGBColor(0, 110, 71),  # 墨绿 - 内置
            Token.String: RGBColor(196, 26, 22),  # 红色 - 字符串
            Token.String.Doc: RGBColor(109, 120, 133),  # 灰 - 文档字符串
            Token.Comment: RGBColor(109, 120, 133),  # 灰 - 注释
            Token.Comment.Single: RGBColor(109, 120, 133),
            Token.Comment.Multiline: RGBColor(109, 120, 133),
            Token.Number: RGBColor(28, 0, 207),  # 靛蓝 - 数字
            Token.Number.Integer: RGBColor(28, 0, 207),
            Token.Number.Float: RGBColor(28, 0, 207),
            Token.Operator: RGBColor(90, 99, 120),  # 灰蓝 - 运算符
            Token.Punctuation: RGBColor(0, 0, 0),  # 黑色 - 标点
        }

        def get_token_color(token_type):
            """递归查找 token 颜色"""
            while token_type:
                if token_type in TOKEN_COLORS:
                    return TOKEN_COLORS[token_type]
                token_type = token_type.parent
            return None

        # 添加语言标签（如果有）
        if language:
            lang_para = doc.add_paragraph()
            lang_para.paragraph_format.space_before = Pt(6)
            lang_para.paragraph_format.space_after = Pt(0)
            lang_para.paragraph_format.left_indent = Cm(0.5)
            lang_run = lang_para.add_run(language.upper())
            lang_run.font.name = "Consolas"
            lang_run.font.size = Pt(8)
            lang_run.font.color.rgb = RGBColor(100, 100, 100)
            lang_run.font.bold = True

        # 添加代码块段落
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.5)
        paragraph.paragraph_format.space_before = Pt(3) if language else Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

        # 添加浅灰色背景
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F7F7F7")
        paragraph._element.pPr.append(shading)

        # 尝试使用 Pygments 进行语法高亮
        if PYGMENTS_AVAILABLE and language:
            try:
                lexer = get_lexer_by_name(language, stripall=False)
            except Exception:
                lexer = TextLexer()

            tokens = list(lex(code, lexer))

            for token_type, token_value in tokens:
                if not token_value:
                    continue
                run = paragraph.add_run(token_value)
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
                run.font.size = Pt(10)

                # 应用颜色
                color = get_token_color(token_type)
                if color:
                    run.font.color.rgb = color

                # 关键字加粗
                if token_type in Token.Keyword:
                    run.font.bold = True
        else:
            # 无语法高亮，纯文本显示
            run = paragraph.add_run(code)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
            run.font.size = Pt(10)

    def add_table(self, doc: Document, table_lines: List[str]):
        """添加表格"""
        if len(table_lines) < 2:
            return

        # 解析表格数据
        rows = []
        for line in table_lines:
            cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
            # 跳过分隔行
            if all(re.fullmatch(r"[-:]+", cell) for cell in cells):
                continue
            rows.append(cells)

        if not rows:
            return

        # 确定列数
        num_cols = max(len(row) for row in rows)

        # 创建表格
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 填充表格
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = row.cells[col_idx]
                    # 清除默认段落
                    cell.paragraphs[0].clear()
                    self.add_formatted_text(cell.paragraphs[0], cell_text)

                    # 设置单元格字体
                    for run in cell.paragraphs[0].runs:
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                        run.font.size = Pt(10)

                    # 表头加粗
                    if row_idx == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True

        # 设置表格列宽度自适应
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_list_to_doc(
        self, doc: Document, items: List[Tuple[int, str]], list_type: str
    ):
        """添加列表"""
        for indent, text in items:
            paragraph = doc.add_paragraph()

            if list_type == "unordered":
                # 无序列表使用项目符号
                paragraph.style = "List Bullet"
            else:
                # 有序列表使用编号
                paragraph.style = "List Number"

            # 设置缩进
            paragraph.paragraph_format.left_indent = Cm(0.5 * (indent + 1))

            # 添加格式化文本
            self.add_formatted_text(paragraph, text)

            # 设置字体
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def add_horizontal_rule(self, doc: Document):
        """添加水平分割线"""
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)

        # 添加底部边框作为分割线
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_blockquote(self, doc: Document, text: str):
        """添加引用块，带有左侧边框和灰色背景"""
        for line in text.split("\n"):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(1.0)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)

            # 添加左侧边框
            pPr = paragraph._element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "24")  # 边框粗细
            left.set(qn("w:space"), "4")  # 边框与文字间距
            left.set(qn("w:color"), "CCCCCC")  # 灰色边框
            pBdr.append(left)
            pPr.append(pBdr)

            # 添加浅灰色背景
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F9F9F9")
            pPr.append(shading)

            # 添加格式化文本
            self.add_formatted_text(paragraph, line)

            # 设置字体为斜体灰色
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
                run.font.color.rgb = RGBColor(85, 85, 85)  # 深灰色文字
                run.italic = True
